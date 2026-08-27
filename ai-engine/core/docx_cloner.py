"""
v19.0: Clone-and-Replace DOCX Generator
========================================
Instead of rebuilding the DOCX from scratch (which loses formatting),
this module CLONES the original DOCX and only replaces the cells
that the AI pipeline has enhanced:

  - B10 (or B7): Department narrative → enhanced_b7
  - D2/E2: Matter summaries → optimized_text per matter

Everything else (colors, bold, numbering, diversity sections, C2 feedback,
client lists, logos, etc.) is preserved EXACTLY as the firm submitted.

Architecture:
  1. Read the original DOCX with python-docx
  2. Walk all tables and identify section markers (B10, D2, E2, etc.)
  3. Match each E2/D2 to its corresponding matter via E1/D1 client name
  4. Replace ONLY the content paragraphs (not the label cells)
  5. Save as new DOCX

Usage:
  from core.docx_cloner import clone_and_replace
  docx_bytes = clone_and_replace(
      original_path="/tmp/uploads/original.docx",
      enhanced_b7="AI-enhanced department narrative...",
      enhanced_matters=[
          {"client": "Grupo Hermes", "optimized_text": "..."},
          {"client": "MEGA DIRECT", "optimized_text": "..."},
      ]
  )
"""

import io
import re
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from copy import deepcopy


# =====================================================
# SECTION DETECTION HEURISTICS
# =====================================================

def _cell_text(cell: _Cell) -> str:
    """Get the full text of a cell, joining all paragraphs."""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _is_b10_label(text: str) -> bool:
    """Check if a cell contains the B10 section label."""
    t = text.lower()
    return ("what is this department best known for" in t or 
            "what is this department" in t)


def _is_matter_summary_label(text: str) -> bool:
    """Check if a cell contains a D2/E2 'Summary of matter' label."""
    t = text.lower()
    return ("summary of matter" in t and "department" in t)


def _is_c2_label(text: str) -> bool:
    """Check if a cell contains the C2 Feedback label."""
    t = text.lower()
    if "b10" in t or "500 word" in t or "best known for" in t:
        return False
    return ("feedback" in t and "other firms" in t) or ("c2" in t and "feedback" in t) or ("feedback on our coverage" in t)


def _is_client_name_label(text: str) -> bool:
    """Check if a cell contains a D1/E1 'Name of client' label."""
    t = text.lower()
    return "name of client" in t


def _normalize_client_name(name: str) -> str:
    """Normalize a client name for fuzzy matching."""
    # Remove descriptors after " - "
    name = name.split(" - ")[0].strip()
    # Remove common corporate and legal suffixes
    for suffix in [", S.A. de C.V.", ", S.A.", " S.A. de C.V.", " S.A.",
                   " S. de R.L.", ", S. de R.L.", " Ltd.", " Ltd",
                   " LLP", " LLC", " N.A.", " Inc.", " Corp.", " Law Firm", " Law"]:
        name = name.replace(suffix, "")
    return name.strip().lower()


def _match_client(client_in_doc: str, enhanced_matters: List[Dict]) -> Optional[Dict]:
    """
    Match a client name from the document to an enhanced matter.
    Uses token intersection and substring matching to handle firm descriptor variations.
    """
    norm_doc = _normalize_client_name(client_in_doc)
    doc_tokens = set(re.findall(r'\w+', norm_doc))
    
    best_match = None
    best_score = 0.0
    
    for matter in enhanced_matters:
        matter_client = matter.get("client", "")
        norm_matter = _normalize_client_name(matter_client)
        matter_tokens = set(re.findall(r'\w+', norm_matter))
        
        # Exact match
        if norm_doc == norm_matter:
            return matter
        
        # Substring / Containment (handles "Kennedys" in "Kennedys and Kennedys Law")
        if norm_matter and (norm_matter in norm_doc or norm_doc in norm_matter):
            return matter
        
        # Token intersection
        intersection = doc_tokens.intersection(matter_tokens)
        meaningful_common = [
            w for w in intersection 
            if w not in {'and', 'the', 'of', 'for', 'group', 'firm', 'solutions', 'limited', 'bank'}
        ]
        if meaningful_common:
            score = len(meaningful_common) / max(len(matter_tokens), 1)
            if score > best_score:
                best_score = score
                best_match = matter
    
    if best_match and best_score >= 0.5:
        return best_match
    
    return None


# =====================================================
# CELL CONTENT REPLACEMENT
# =====================================================

def strip_system_instructions(text: str) -> str:
    """v24.3: Prevent system reasoning or internal directives from leaking into client-facing DOCX."""
    if not text:
        return ""
    instruction_patterns = [
        r'Recover the complete evidentiary record.*',
        r'Test practice trajectory only after recovery.*',
        r'Select the hero matter only after all candidates can be assessed.*',
        r'\[INTERNAL DIRECTIVE\].*',
        r'\[SYSTEM NOTE\].*',
        r'\[RETRY DIRECTIVE\].*',
        r'Address any feedback.*word count limit\)?',
        r'Please include:.*word count limit\)?'
    ]
    cleaned = text
    for pattern in instruction_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE | re.DOTALL).strip()
    return cleaned


def _replace_cell_content(cell: _Cell, new_text: str, preserve_first_paragraph_format: bool = True):
    """
    Replace the content of a cell while preserving paragraph formatting.
    
    Strategy:
    1. Split new_text by double newlines into paragraphs
    2. For the first paragraph, preserve the original paragraph's run formatting
    3. For subsequent paragraphs, clone formatting from the first paragraph
    4. Remove any excess original paragraphs
    """
    new_text = strip_system_instructions(new_text)
    new_paragraphs = [p.strip() for p in new_text.split("\n\n") if p.strip()]
    
    if not new_paragraphs:
        return
    
    original_paragraphs = cell.paragraphs
    
    if not original_paragraphs:
        return
    
    # Get reference formatting from the first content paragraph
    ref_paragraph = original_paragraphs[0]
    ref_format = None
    if ref_paragraph.runs:
        ref_run = ref_paragraph.runs[0]
        ref_format = {
            "font_name": ref_run.font.name,
            "font_size": ref_run.font.size,
            "bold": ref_run.font.bold,
            "italic": ref_run.font.italic,
            "color": ref_run.font.color.rgb if ref_run.font.color and ref_run.font.color.rgb else None,
        }
    
    # Clear all existing paragraphs
    for i in range(len(original_paragraphs) - 1, -1, -1):
        p = original_paragraphs[i]
        p_element = p._element
        p_element.getparent().remove(p_element)
    
    # Add new paragraphs
    for i, para_text in enumerate(new_paragraphs):
        if i == 0:
            # Re-add the first paragraph (we deleted it)
            from docx.oxml.ns import qn
            new_p = cell._element.makeelement(qn('w:p'), {})
            cell._element.append(new_p)
            p = Paragraph(new_p, cell)
        else:
            p = cell.add_paragraph()
        
        # Check for bold markers: **text** → bold
        segments = _parse_bold_segments(para_text)
        
        for text, is_bold in segments:
            run = p.add_run(text)
            if ref_format:
                if ref_format["font_name"]:
                    run.font.name = ref_format["font_name"]
                if ref_format["font_size"]:
                    run.font.size = ref_format["font_size"]
                if ref_format["color"]:
                    run.font.color.rgb = ref_format["color"]
                # Bold: either from the ** markers or from original formatting
                run.font.bold = is_bold or (ref_format["bold"] if not segments[0][1] else False)
                if ref_format["italic"]:
                    run.font.italic = ref_format["italic"]


def _parse_bold_segments(text: str) -> List[Tuple[str, bool]]:
    """
    Parse text with **bold markers** into segments.
    Returns list of (text, is_bold) tuples.
    """
    segments = []
    pattern = re.compile(r'\*\*(.*?)\*\*')
    last_end = 0
    
    for match in pattern.finditer(text):
        # Text before the bold marker
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], False))
        # Bold text
        segments.append((match.group(1), True))
        last_end = match.end()
    
    # Remaining text after last bold marker
    if last_end < len(text):
        segments.append((text[last_end:], False))
    
    if not segments:
        segments.append((text, False))
    
    return segments


# =====================================================
# TABLE STRUCTURE ANALYSIS
# =====================================================

def _find_data_cell_in_table(table: Table, label_check_fn) -> Optional[Tuple[int, int]]:
    """
    Find a label cell in a table and return the position of the DATA cell below it.
    Chambers template pattern: label row → data row (in the next row, same column).
    Returns (row_index, col_index) of the data cell, or None.
    """
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            text = _cell_text(cell)
            if label_check_fn(text):
                # The data cell is typically in the NEXT row, same column
                if row_idx + 1 < len(table.rows):
                    return (row_idx + 1, col_idx)
    return None


def _find_client_name_in_table(table: Table) -> str:
    """
    Find the E1/D1 client name in a matter table.
    Pattern: "Name of client" label → data in next row.
    """
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            text = _cell_text(cell)
            if _is_client_name_label(text):
                if row_idx + 1 < len(table.rows):
                    return _cell_text(table.rows[row_idx + 1].cells[0])
    return ""


def _append_matter_table(doc: Document, matter_idx: int, matter: Dict, is_confidential: bool = False):
    """Append a standard Chambers D/E matter table to the document."""
    from docx.shared import Pt
    prefix = 'E' if is_confidential else 'D'
    sec_label = 'CONFIDENTIAL' if is_confidential else 'PUBLISHABLE'
    
    # Add page break before matter
    doc.add_page_break()
    p_header = doc.add_paragraph()
    p_run = p_header.add_run(f"Matter {matter_idx} ({sec_label} WORK HIGHLIGHT)")
    p_run.bold = True
    p_run.font.size = Pt(12)
    
    table = doc.add_table(rows=0, cols=1)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    
    # 1. Client
    r1 = table.add_row()
    r1.cells[0].text = f"{prefix}1 Name of client (including country of origin and website URL):"
    if r1.cells[0].paragraphs and r1.cells[0].paragraphs[0].runs:
        r1.cells[0].paragraphs[0].runs[0].bold = True
    r2 = table.add_row()
    r2.cells[0].text = matter.get("client", "")
    
    # 2. Summary
    r3 = table.add_row()
    r3.cells[0].text = f"{prefix}2 Summary of matter and your department's involvement:"
    if r3.cells[0].paragraphs and r3.cells[0].paragraphs[0].runs:
        r3.cells[0].paragraphs[0].runs[0].bold = True
    r4 = table.add_row()
    _replace_cell_content(r4.cells[0], matter.get("optimized_text") or matter.get("summary", ""))
    
    # 3. Value
    r5 = table.add_row()
    r5.cells[0].text = f"{prefix}3 Value of deal / matter (if applicable):"
    if r5.cells[0].paragraphs and r5.cells[0].paragraphs[0].runs:
        r5.cells[0].paragraphs[0].runs[0].bold = True
    r6 = table.add_row()
    r6.cells[0].text = str(matter.get("value", "") or "N/A")
    
    # 4. Lead partner
    r7 = table.add_row()
    r7.cells[0].text = f"{prefix}4 Lead partner / lawyers involved:"
    if r7.cells[0].paragraphs and r7.cells[0].paragraphs[0].runs:
        r7.cells[0].paragraphs[0].runs[0].bold = True
    r8 = table.add_row()
    r8.cells[0].text = str(matter.get("lead_partner", "") or matter.get("leadPartner", "") or "")


# =====================================================
# MAIN CLONE-AND-REPLACE FUNCTION
# =====================================================

def clone_and_replace(
    original_path: str,
    enhanced_b7: str = "",
    enhanced_matters: Optional[List[Dict]] = None,
    enhanced_c2: str = "",
) -> bytes:
    """
    Clone the original DOCX and replace only B10 + C2 + D2/E2 cells.
    If the original document lacks C2 or D2/E2 tables, appends them cleanly.
    
    Args:
        original_path: Path to the original DOCX file uploaded by the firm
        enhanced_b7: AI-enhanced department narrative (replaces B10 cell content)
        enhanced_matters: List of dicts with {"client": str, "optimized_text": str, ...}
                         Each replaces the corresponding D2/E2 matter summary
        enhanced_c2: AI-enhanced C2 feedback narrative (replaces C2 cell content)
    
    Returns:
        bytes: The modified DOCX file as bytes
    """
    from docx.shared import Pt
    if enhanced_matters is None:
        enhanced_matters = []
    
    print(f"[DOCX CLONER] Loading original document: {original_path}")
    doc = Document(original_path)
    
    b7_replaced = False
    c2_replaced = False
    matters_replaced = 0
    matters_skipped = []
    
    # Track which enhanced matters have been used
    used_matters = set()
    
    # Collect all tables across the document tree (including nested <w:sdt> controls)
    all_doc_tables = [Table(elem, doc) for elem in doc.element.body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')]
    
    for table_idx, table in enumerate(all_doc_tables):
        # ─── CHECK FOR B10 SECTION ───
        if enhanced_b7 and not b7_replaced:
            data_pos = _find_data_cell_in_table(table, _is_b10_label)
            if data_pos:
                row_idx, col_idx = data_pos
                data_cell = table.rows[row_idx].cells[col_idx]
                original_text = _cell_text(data_cell)
                
                print(f"[DOCX CLONER] Found B10 at table {table_idx}, row {row_idx}")
                print(f"[DOCX CLONER]   Original B10: {len(original_text.split())} words")
                print(f"[DOCX CLONER]   Enhanced B7:  {len(enhanced_b7.split())} words")
                
                _replace_cell_content(data_cell, enhanced_b7)
                b7_replaced = True
                continue
        
        # ─── CHECK FOR C2 SECTION ───
        if enhanced_c2 and not c2_replaced:
            data_pos = _find_data_cell_in_table(table, _is_c2_label)
            if data_pos:
                row_idx, col_idx = data_pos
                data_cell = table.rows[row_idx].cells[col_idx]
                original_text = _cell_text(data_cell)
                
                print(f"[DOCX CLONER] Found C2 at table {table_idx}, row {row_idx}")
                print(f"[DOCX CLONER]   Original C2: {len(original_text.split())} words")
                print(f"[DOCX CLONER]   Enhanced C2: {len(enhanced_c2.split())} words")
                
                _replace_cell_content(data_cell, enhanced_c2)
                c2_replaced = True
                continue
        
    # ─── v25.2 STRATEGIC PHYSICAL TABLE RE-ORDERING ───
    # Ensure Hero Matter is written to Physical Table 1 in Section D
    if enhanced_matters:
        # Collect all physical matter tables in Section D and Section E
        sec_d_tables = []
        sec_e_tables = []
        
        for table in all_doc_tables:
            table_text = "\n".join(_cell_text(c) for r in table.rows for c in r.cells).lower()
            if "d1 name of client" in table_text or ("d2 summary" in table_text and "d1" in table_text):
                sec_d_tables.append(table)
            elif "e1 name of client" in table_text or ("e2 summary" in table_text and "e1" in table_text):
                sec_e_tables.append(table)
        
        pub_matters = [
            m for m in enhanced_matters 
            if not (m.get("is_confidential") or m.get("publish_status") == "confidential")
        ]
        conf_matters = [
            m for m in enhanced_matters 
            if (m.get("is_confidential") or m.get("publish_status") == "confidential")
        ]

        print(f"[DOCX CLONER 🎯] Physical tables found: {len(sec_d_tables)} in Section D, {len(sec_e_tables)} in Section E")
        print(f"[DOCX CLONER 🎯] Enhanced matters to write: {len(pub_matters)} Publishable, {len(conf_matters)} Confidential")

        # 1. Populate Section D tables in exact strategic order (Hero Matter #1 in Table 1)
        for idx, m in enumerate(pub_matters):
            client_name = m.get("client", "")
            opt_text = m.get("optimized_text") or m.get("summary", "")
            val_text = str(m.get("value", "") or "N/A")
            partner_text = str(m.get("lead_partner", "") or m.get("leadPartner", "") or "")

            if idx < len(sec_d_tables):
                t = sec_d_tables[idx]
                for i, r in enumerate(t.rows):
                    c0 = r.cells[0].text.strip().lower()
                    if "d1 name" in c0 and i + 1 < len(t.rows):
                        _replace_cell_content(t.rows[i+1].cells[0], client_name)
                    elif "d2 summary" in c0 and i + 1 < len(t.rows) and opt_text:
                        _replace_cell_content(t.rows[i+1].cells[0], opt_text)
                    elif "d3 value" in c0 and i + 1 < len(t.rows):
                        _replace_cell_content(t.rows[i+1].cells[0], val_text)
                    elif "d4 lead partner" in c0 and i + 1 < len(t.rows):
                        _replace_cell_content(t.rows[i+1].cells[0], partner_text)
                
                client_key = _normalize_client_name(client_name)
                used_matters.add(client_key)
                matters_replaced += 1
                if idx == 0:
                    print(f"[DOCX CLONER 🏆] HERO MATTER SUCCESSFULLY WRITTEN TO SECTION D TABLE 1: {client_name}")

        # 2. Populate Section E tables in exact strategic order
        for idx, m in enumerate(conf_matters):
            client_name = m.get("client", "")
            opt_text = m.get("optimized_text") or m.get("summary", "")
            val_text = str(m.get("value", "") or "N/A")
            partner_text = str(m.get("lead_partner", "") or m.get("leadPartner", "") or "")

            if idx < len(sec_e_tables):
                t = sec_e_tables[idx]
                for i, r in enumerate(t.rows):
                    c0 = r.cells[0].text.strip().lower()
                    if "e1 name" in c0 and i + 1 < len(t.rows):
                        _replace_cell_content(t.rows[i+1].cells[0], client_name)
                    elif "e2 summary" in c0 and i + 1 < len(t.rows) and opt_text:
                        _replace_cell_content(t.rows[i+1].cells[0], opt_text)
                    elif "e3 value" in c0 and i + 1 < len(t.rows):
                        _replace_cell_content(t.rows[i+1].cells[0], val_text)
                    elif "e4 lead partner" in c0 and i + 1 < len(t.rows):
                        _replace_cell_content(t.rows[i+1].cells[0], partner_text)

                client_key = _normalize_client_name(client_name)
                used_matters.add(client_key)
                matters_replaced += 1
    
    # ─── v23.0 IN-PLACE INSERTION FOR MISSING C2 & MATTERS ───
    # Find anchor points in original document
    sec_d_para = None
    sec_e_para = None
    c1_table = None
    
    for p in doc.paragraphs:
        t = p.text.strip().lower()
        if 'd. publishable information' in t:
            sec_d_para = p
        elif 'e. confidential information' in t:
            sec_e_para = p
    
    for t in all_doc_tables:
        for r in t.rows:
            for c in r.cells:
                if 'barrister' in c.text.lower() or 'c1' in c.text.lower():
                    c1_table = t
                    break
    
    # 1. Insert C2 Table at Section C anchor if not found in existing tables
    if enhanced_c2 and not c2_replaced:
        print(f"[DOCX CLONER] Inserting C2 Feedback table in Section C ({len(enhanced_c2.split())} words)...")
        c2_table = doc.add_table(rows=0, cols=1)
        try:
            c2_table.style = 'Table Grid'
        except Exception:
            pass
        
        r1 = c2_table.add_row()
        r1.cells[0].text = "C2 Feedback on our coverage of this practice area (Optional):"
        if r1.cells[0].paragraphs and r1.cells[0].paragraphs[0].runs:
            r1.cells[0].paragraphs[0].runs[0].bold = True
        r2 = c2_table.add_row()
        _replace_cell_content(r2.cells[0], enhanced_c2)
        
        if c1_table:
            c1_table._element.addnext(c2_table._element)
        c2_replaced = True
    
    # 2. Insert Unmatched Matters at Section D and Section E anchors
    unmatched_matters = [
        m for m in enhanced_matters 
        if _normalize_client_name(m.get("client", "")) not in used_matters
    ]
    if unmatched_matters:
        print(f"[DOCX CLONER] Inserting {len(unmatched_matters)} matter tables in Sections D and E...")
        
        pub_matters = [
            m for m in unmatched_matters 
            if not (m.get("is_confidential") or m.get("publish_status") == "confidential")
        ]
        conf_matters = [
            m for m in unmatched_matters 
            if (m.get("is_confidential") or m.get("publish_status") == "confidential")
        ]
        
        # Insert publishable matters inside Section D
        if sec_d_para and pub_matters:
            for m in reversed(pub_matters):
                t_pub = doc.add_table(rows=0, cols=1)
                try:
                    t_pub.style = 'Table Grid'
                except Exception:
                    pass
                r1 = t_pub.add_row()
                r1.cells[0].text = "D1 Name of client (including country of origin and website URL):"
                if r1.cells[0].paragraphs and r1.cells[0].paragraphs[0].runs:
                    r1.cells[0].paragraphs[0].runs[0].bold = True
                r2 = t_pub.add_row()
                r2.cells[0].text = m.get("client", "")
                r3 = t_pub.add_row()
                r3.cells[0].text = "D2 Summary of matter and your department's involvement:"
                if r3.cells[0].paragraphs and r3.cells[0].paragraphs[0].runs:
                    r3.cells[0].paragraphs[0].runs[0].bold = True
                r4 = t_pub.add_row()
                _replace_cell_content(r4.cells[0], m.get("optimized_text") or m.get("summary", ""))
                r5 = t_pub.add_row()
                r5.cells[0].text = "D3 Value of deal / matter (if applicable):"
                if r5.cells[0].paragraphs and r5.cells[0].paragraphs[0].runs:
                    r5.cells[0].paragraphs[0].runs[0].bold = True
                r6 = t_pub.add_row()
                r6.cells[0].text = str(m.get("value", "") or "N/A")
                r7 = t_pub.add_row()
                r7.cells[0].text = "D4 Lead partner / lawyers involved:"
                if r7.cells[0].paragraphs and r7.cells[0].paragraphs[0].runs:
                    r7.cells[0].paragraphs[0].runs[0].bold = True
                r8 = t_pub.add_row()
                r8.cells[0].text = str(m.get("lead_partner", "") or m.get("leadPartner", "") or "")
                
                sec_d_para._element.addnext(t_pub._element)
                matters_replaced += 1
                used_matters.add(_normalize_client_name(m.get("client", "")))
        
        # Insert confidential matters inside Section E
        if sec_e_para and conf_matters:
            for m in reversed(conf_matters):
                t_conf = doc.add_table(rows=0, cols=1)
                try:
                    t_conf.style = 'Table Grid'
                except Exception:
                    pass
                r1 = t_conf.add_row()
                r1.cells[0].text = "E1 Name of client (including country of origin and website URL):"
                if r1.cells[0].paragraphs and r1.cells[0].paragraphs[0].runs:
                    r1.cells[0].paragraphs[0].runs[0].bold = True
                r2 = t_conf.add_row()
                r2.cells[0].text = m.get("client", "")
                r3 = t_conf.add_row()
                r3.cells[0].text = "E2 Summary of matter and your department's involvement:"
                if r3.cells[0].paragraphs and r3.cells[0].paragraphs[0].runs:
                    r3.cells[0].paragraphs[0].runs[0].bold = True
                r4 = t_conf.add_row()
                _replace_cell_content(r4.cells[0], m.get("optimized_text") or m.get("summary", ""))
                r5 = t_conf.add_row()
                r5.cells[0].text = "E3 Value of deal / matter (if applicable):"
                if r5.cells[0].paragraphs and r5.cells[0].paragraphs[0].runs:
                    r5.cells[0].paragraphs[0].runs[0].bold = True
                r6 = t_conf.add_row()
                r6.cells[0].text = str(m.get("value", "") or "N/A")
                r7 = t_conf.add_row()
                r7.cells[0].text = "E4 Lead partner / lawyers involved:"
                if r7.cells[0].paragraphs and r7.cells[0].paragraphs[0].runs:
                    r7.cells[0].paragraphs[0].runs[0].bold = True
                r8 = t_conf.add_row()
                r8.cells[0].text = str(m.get("lead_partner", "") or m.get("leadPartner", "") or "")
                
                sec_e_para._element.addnext(t_conf._element)
                matters_replaced += 1
                used_matters.add(_normalize_client_name(m.get("client", "")))
    
    # ─── SUMMARY ───
    print(f"\n[DOCX CLONER] ════════════════════════════════════════")
    print(f"[DOCX CLONER] Clone-and-Replace complete:")
    print(f"[DOCX CLONER]   B10 replaced: {'✅ Yes' if b7_replaced else '❌ No (not found or no enhanced_b7)'}")
    print(f"[DOCX CLONER]   C2 replaced/inserted: {'✅ Yes' if c2_replaced else '❌ No'}")
    print(f"[DOCX CLONER]   Matters present: {matters_replaced}/{len(enhanced_matters)}")
    if matters_skipped:
        print(f"[DOCX CLONER]   Skipped (no match): {', '.join(s[:40] for s in matters_skipped)}")
    print(f"[DOCX CLONER] ════════════════════════════════════════\n")
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def clone_and_replace_from_state(
    file_path: str,
    enhanced_b7: str,
    matters: List[Dict],
    enhanced_c2: str = "",
    hero_matter: str = "",
    matter_order: List[str] = None,
) -> Optional[bytes]:
    """
    Convenience wrapper that extracts enhanced_matters from the pipeline state format.
    
    Handles both local file paths and URLs (Supabase Storage).
    When file_path is a URL, downloads the DOCX to a temp file first.
    
    Args:
        file_path: Path or URL to original DOCX
        enhanced_b7: AI-enhanced B7 narrative
        matters: List of matter dicts from the pipeline (with client + optimized_text fields)
        enhanced_c2: AI-enhanced C2 feedback narrative
        hero_matter: Name of hero matter to force to position #1
        matter_order: Desired order of client names from Blueprint
    
    Returns:
        bytes or None if file_path is invalid or not a DOCX
    """
    import os
    import tempfile
    import urllib.request
    from urllib.parse import urlparse
    
    if not file_path or not isinstance(file_path, str) or not file_path.strip():
        # Fallback to official native Chambers template on disk
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        fallback_template = os.path.join(base_dir, "templates", "chambers_template.docx")
        if os.path.exists(fallback_template):
            print(f"[DOCX CLONER] No input DOCX file provided — using official native disk template: {fallback_template}")
            file_path = fallback_template
        else:
            print("[DOCX CLONER] Skipping — no file_path provided and no fallback template found")
            return None
    
    # Check if it's a URL or local path
    is_url = file_path.startswith('http://') or file_path.startswith('https://')
    
    # Validate extension (handle URL query params)
    parsed_path = urlparse(file_path).path if is_url else file_path
    if not parsed_path.lower().endswith('.docx'):
        print(f"[DOCX CLONER] Skipping — not a DOCX file: {file_path[:100]}")
        return None
    
    # Build enhanced_matters list from pipeline state (preserving full metadata)
    enhanced_matters = []
    for m in matters:
        optimized = m.get("optimized_text") or m.get("optimizedText") or m.get("summary")
        client = m.get("client", "")
        if client:
            matter_entry = dict(m)
            if optimized:
                matter_entry["optimized_text"] = optimized
            enhanced_matters.append(matter_entry)

    # v25.0: BLUEPRINT ORDER ENFORCEMENT — Force Hero Matter to position #1
    if hero_matter:
        norm_hero = _normalize_client_name(hero_matter)
        hero_matches = []
        non_hero = []
        for m in enhanced_matters:
            norm_c = _normalize_client_name(m.get("client", ""))
            title = m.get("title", "").lower()
            if norm_hero in norm_c or norm_c in norm_hero or hero_matter.lower() in title or hero_matter.lower() in m.get("client", "").lower():
                hero_matches.append(m)
            else:
                non_hero.append(m)
        if hero_matches:
            enhanced_matters = hero_matches + non_hero
            print(f"[DOCX CLONER 🎯] BLUEPRINT EXECUTION: Hero Matter '{hero_matches[0].get('client')}' forced to Position #1")
    
    if not enhanced_b7 and not enhanced_matters and not enhanced_c2:
        print("[DOCX CLONER] Skipping — no enhanced content to replace")
        return None
    
    # Download from URL if needed
    local_path = file_path
    temp_file = None
    
    if is_url:
        try:
            print(f"[DOCX CLONER] Downloading original DOCX from URL...")
            fd, local_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            req = urllib.request.Request(file_path, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=30) as response, open(local_path, 'wb') as out_file:
                out_file.write(response.read())
            temp_file = local_path
            print(f"[DOCX CLONER] Downloaded to temp file: {local_path} ({os.path.getsize(local_path)} bytes)")
        except Exception as dl_err:
            print(f"[DOCX CLONER] Failed to download DOCX: {dl_err}")
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)
            return None
    elif not os.path.exists(file_path):
        print(f"[DOCX CLONER] Skipping — local file not found: {file_path}")
        return None
    
    try:
        return clone_and_replace(
            original_path=local_path,
            enhanced_b7=enhanced_b7,
            enhanced_matters=enhanced_matters,
            enhanced_c2=enhanced_c2,
        )
    finally:
        # Cleanup temp file
        if temp_file and os.path.exists(temp_file):
            os.remove(temp_file)
            print(f"[DOCX CLONER] Cleaned up temp file")

