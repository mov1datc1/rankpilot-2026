from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_docx_report(structured_data: dict, output_filename: str, doc_type: str = 'audit') -> str:
    """
    Generates a professional DOCX report using python-docx.
    Supports doc_type: 'audit' (internal) or 'submission' (Chambers Template).
    v14.0: Added Pipeline Manifest page for auditability (Rule 71).
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'chambers_template.docx')
    
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
        title = doc.add_heading(level=0)
        title_run = title.add_run('RANKPILOT OFFICIAL SUBMISSION')
        title_run.font.color.rgb = RGBColor(26, 35, 126) # Brand Navy
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Extraer los datos estratégicos de la sesión
    chambers_data = structured_data.get('chambersData', {})
    
    firm_name = structured_data.get('firm_metadata', {}).get('firm_name', '') or chambers_data.get('firmName', 'Professional Law Firm')
    practice_area = structured_data.get('firm_metadata', {}).get('practice_area', '') or chambers_data.get('practice', 'General Practice')
    analysis = chambers_data.get('analysis', {})
    context = chambers_data.get('strategicContext', {})
    letter = analysis.get('audit_letter', {})

    if doc_type == 'audit':
        # =====================================================
        # v14.0 TRUST LAYER — Rule 71: Pipeline Manifest Page
        # First page of every audit report. Answers:
        # - What did the system read?
        # - How many matters found vs. source?
        # - What context/RAG files were loaded?
        # =====================================================
        manifest = structured_data.get('pipeline_manifest', {}) or chambers_data.get('pipeline_manifest', {})
        if manifest and manifest.get('document'):
            doc.add_page_break()
            manifest_title = doc.add_heading('Pipeline Manifest — Trust Layer', level=1)
            manifest_title.runs[0].font.color.rgb = RGBColor(26, 35, 126)
            
            doc_info = manifest.get('document', {})
            p = doc.add_paragraph()
            p.add_run('What the system read:\n').bold = True
            p.add_run(f"File: {doc_info.get('file_name', 'Unknown')}\n")
            p.add_run(f"Hash: {doc_info.get('file_hash', 'N/A')}\n")
            p.add_run(f"Words: {doc_info.get('word_count', 0)} | Paragraphs: {doc_info.get('paragraph_count', 0)} | Tables: {doc_info.get('table_count', 0)}\n")
            
            source_matters = doc_info.get('source_matters', {})
            extraction = manifest.get('extraction', {})
            
            p2 = doc.add_paragraph()
            p2.add_run('Matter Verification:\n').bold = True
            p2.add_run(f"Source document matters: {source_matters.get('total', 'N/A')}")
            if source_matters.get('publishable', 0) or source_matters.get('confidential', 0):
                p2.add_run(f" (publishable: {source_matters.get('publishable', 0)}, confidential: {source_matters.get('confidential', 0)})")
            p2.add_run(f"\nExtracted by AI: {extraction.get('extracted_matter_count', 'N/A')}\n")
            
            if extraction.get('loss_count', 0) > 0:
                p_warn = doc.add_paragraph()
                warn_run = p_warn.add_run(f"⚠️ MATTER LOSS DETECTED: {extraction['loss_count']} matters lost ({extraction.get('loss_percentage', 0)}%)")
                warn_run.bold = True
                warn_run.font.color.rgb = RGBColor(200, 0, 0)
            elif extraction.get('match'):
                p_ok = doc.add_paragraph()
                ok_run = p_ok.add_run("✅ Matter count VERIFIED — extraction matches source document")
                ok_run.bold = True
                ok_run.font.color.rgb = RGBColor(0, 128, 0)
            
            if source_matters.get('matter_labels'):
                p3 = doc.add_paragraph()
                p3.add_run('\nSource matter labels:\n').bold = True
                for label in source_matters['matter_labels']:
                    doc.add_paragraph(label, style='List Bullet')
            
            if extraction.get('extracted_titles'):
                p4 = doc.add_paragraph()
                p4.add_run('\nExtracted matter titles:\n').bold = True
                for title_text in extraction['extracted_titles']:
                    doc.add_paragraph(title_text, style='List Bullet')
            
            rag_files = manifest.get('rag_files_loaded', [])
            if rag_files:
                p5 = doc.add_paragraph()
                p5.add_run('\nRAG Knowledge Files Loaded:\n').bold = True
                for f_name in rag_files:
                    doc.add_paragraph(f_name, style='List Bullet')
            
            doc.add_paragraph(f"\nTimestamp: {manifest.get('timestamp', 'N/A')}")
            
            # =====================================================
            # v14.1 PRE-FLIGHT GATE — Rule 74: Validation Results
            # Shows the 5-point check results in the DOCX
            # =====================================================
            pre_flight = manifest.get('pre_flight', {})
            if pre_flight and pre_flight.get('checks'):
                doc.add_paragraph()
                pf_heading = doc.add_heading('Pre-Flight Gate — Validation Results', level=2)
                pf_heading.runs[0].font.color.rgb = RGBColor(26, 35, 126)
                
                gate_status = "✅ ALL CHECKS PASSED" if pre_flight.get('passed', True) else "❌ PIPELINE HALTED"
                p_gate = doc.add_paragraph()
                gate_run = p_gate.add_run(f"Gate Status: {gate_status}")
                gate_run.bold = True
                gate_run.font.color.rgb = RGBColor(0, 128, 0) if pre_flight.get('passed', True) else RGBColor(200, 0, 0)
                
                for check in pre_flight['checks']:
                    status_icon = {"PASS": "✅", "WARN": "⚠️", "FAIL": "❌", "SKIP": "⏭️"}.get(check.get('status', ''), '•')
                    p_check = doc.add_paragraph(style='List Bullet')
                    status_run = p_check.add_run(f"{status_icon} {check.get('name', '')}: ")
                    status_run.bold = True
                    if check.get('status') == 'FAIL':
                        status_run.font.color.rgb = RGBColor(200, 0, 0)
                    elif check.get('status') == 'WARN':
                        status_run.font.color.rgb = RGBColor(200, 150, 0)
                    p_check.add_run(check.get('detail', ''))
                
                if pre_flight.get('warnings'):
                    p_warnings = doc.add_paragraph()
                    p_warnings.add_run(f"\nWarnings ({len(pre_flight['warnings'])}):").bold = True
                    for w in pre_flight['warnings']:
                        doc.add_paragraph(f"⚠️ {w}", style='List Bullet')
                
                if pre_flight.get('errors'):
                    p_errors = doc.add_paragraph()
                    err_run = p_errors.add_run(f"\nCritical Errors ({len(pre_flight['errors'])}):")
                    err_run.bold = True
                    err_run.font.color.rgb = RGBColor(200, 0, 0)
                    for e in pre_flight['errors']:
                        doc.add_paragraph(f"❌ {e}", style='List Bullet')
                
                # Template Detection Results
                template = manifest.get('template_detection', {})
                if template and template.get('detected_directory', 'Unknown') != 'Unknown':
                    p_template = doc.add_paragraph()
                    p_template.add_run('\nAuto-Detected Document Identity:\n').bold = True
                    p_template.add_run(f"Directory: {template.get('detected_directory', 'Unknown')} ({template.get('confidence', 'low')} confidence)\n")
                    if template.get('detected_firm_name'):
                        p_template.add_run(f"Firm: {template['detected_firm_name']}\n")
                    if template.get('detected_practice_area'):
                        p_template.add_run(f"Practice Area: {template['detected_practice_area']}\n")
                    if template.get('detected_jurisdiction'):
                        p_template.add_run(f"Jurisdiction: {template['detected_jurisdiction']}\n")
                
                # Ranking Evidence
                ranking = manifest.get('ranking_evidence', {})
                if ranking and ranking.get('has_ranking_evidence'):
                    p_ranking = doc.add_paragraph()
                    rank_run = p_ranking.add_run('\n⚠️ Ranking Evidence Detected:\n')
                    rank_run.bold = True
                    rank_run.font.color.rgb = RGBColor(200, 150, 0)
                    p_ranking.add_run(f"Type: {ranking.get('evidence_type', 'unknown')}\n")
                    p_ranking.add_run(f"Evidence: \"{ranking.get('evidence_text', '')}\"")
                    if ranking.get('detected_band'):
                        p_ranking.add_run(f"\nDetected Band: {ranking['detected_band']}")

        doc.add_page_break()
        audit_title = doc.add_heading('Strategic Audit Letter', level=1)
        audit_title.runs[0].font.color.rgb = RGBColor(26, 35, 126)
        
        # Metadatos del Audit
        p_meta = doc.add_paragraph()
        p_meta.add_run('To: ').bold = True
        p_meta.add_run('The Board of Directors at the Firm\n')
        p_meta.add_run('From: ').bold = True
        p_meta.add_run('RankPilot Consulting\n')

        # Executive Summary
        if analysis.get('summary'):
            doc.add_heading('Executive Summary', level=2)
            p_sum = doc.add_paragraph(str(analysis.get('summary')))
            p_sum.italic = True

        # State of Play
        if letter.get('the_state_of_play'):
            doc.add_heading('The State of Play', level=2)
            doc.add_paragraph(str(letter.get('the_state_of_play')))

        # Reality Check
        reality = letter.get('the_reality_check', [])
        if reality and isinstance(reality, list):
            doc.add_heading('The Reality Check', level=2)
            doc.add_paragraph('Editorial observations on the submission\'s competitive positioning:')
            for item in reality:
                doc.add_paragraph(str(item), style='List Bullet')

        # Path to Dominance
        path = letter.get('the_path_to_dominance', [])
        if path and isinstance(path, list):
            doc.add_heading('The Path to Dominance', level=2)
            for idx, step in enumerate(path, 1):
                title = step.get('title', 'Strategic Step') if isinstance(step, dict) else 'Strategic Step'
                desc = step.get('description', str(step)) if isinstance(step, dict) else str(step)
                p_step = doc.add_paragraph()
                p_step.add_run(f"STEP {idx}: {title}").bold = True
                doc.add_paragraph(desc)
                
    elif doc_type == 'submission':
        # --- OFFICIAL CHAMBERS TEMPLATE ---
        doc.add_paragraph("SUBMISSION FORM", style='Heading 1')
        doc.add_paragraph("Please do not alter this submission template. If a question does not apply to you, please leave it blank.")
        doc.add_paragraph("If something is confidential, mark it as such throughout.")
        
        doc.add_heading('A. PRELIMINARY INFORMATION', level=1)
        p_a = doc.add_paragraph()
        p_a.add_run('Firm Name: ').bold = True
        p_a.add_run(f"{firm_name}\n")
        p_a.add_run('Practice Area: ').bold = True
        p_a.add_run(f"{practice_area}\n")
        if chambers_data.get('jurisdiction'):
            p_a.add_run('Jurisdiction: ').bold = True
            p_a.add_run(f"{chambers_data.get('jurisdiction')}\n")
            
        doc.add_heading('B. DEPARTMENT INFORMATION', level=1)
        if chambers_data.get('departmentDesc'):
            doc.add_paragraph().add_run('Department Description:').bold = True
            doc.add_paragraph(str(chambers_data.get('departmentDesc')))
        if chambers_data.get('specialties'):
            doc.add_paragraph().add_run('Key Specialties:').bold = True
            doc.add_paragraph(str(chambers_data.get('specialties')))
            
        doc.add_heading('C. FEEDBACK', level=1)
        if chambers_data.get('feedback'):
            doc.add_paragraph(str(chambers_data.get('feedback')))
        if chambers_data.get('differentiators'):
            doc.add_paragraph().add_run('Differentiators:').bold = True
            doc.add_paragraph(str(chambers_data.get('differentiators')))
            
        doc.add_heading('WORK HIGHLIGHTS AND CLIENTS', level=1)
        doc.add_paragraph("Provide details of up-to a total of 20 work highlights for this area.")
        
        matters = structured_data.get('matters', [])
        
        # v11.0: Split matters by publish status
        publishable_matters = [m for m in matters if m.get('publish_status', 'publishable') == 'publishable' and not m.get('is_confidential', False)]
        confidential_matters = [m for m in matters if m.get('publish_status', 'publishable') != 'publishable' or m.get('is_confidential', False)]
        
        # Section D: Publishable matters
        doc.add_heading('D. PUBLISHABLE INFORMATION', level=2)
        
        if publishable_matters:
            for idx, matter in enumerate(publishable_matters, 1):
                doc.add_heading(f"Matter {idx}: {matter.get('name', '') or matter.get('title', 'Untitled')}", level=3)
                
                p = doc.add_paragraph()
                p.add_run('Client: ').bold = True
                p.add_run(f"{matter.get('client', 'N/A')}\n")
                
                p.add_run('Value: ').bold = True
                p.add_run(f"{matter.get('value', '') or matter.get('matter_value', 'N/A')}\n")
                
                p.add_run('Lead Partner: ').bold = True
                p.add_run(f"{matter.get('lead_partner') or matter.get('leadPartner', 'N/A')}\n")
                
                doc.add_paragraph().add_run('Matter Summary (Publishable):').bold = True
                doc.add_paragraph(matter.get('optimizedText') or matter.get('optimized_text') or matter.get('description') or matter.get('rawNotes') or '')
                doc.add_paragraph("_" * 50)
        else:
            doc.add_paragraph("No publishable matters available. Review Section E for confidential matters that may be reclassified.")
        
        # Section E: Non-publishable / Confidential matters (v11.0)
        if confidential_matters:
            doc.add_heading('E. NON-PUBLISHABLE / CONFIDENTIAL MATTERS', level=2)
            doc.add_paragraph("The following matters are marked as confidential or non-publishable. They will NOT appear in the published directory entry.")
            
            for idx, matter in enumerate(confidential_matters, 1):
                doc.add_heading(f"Confidential Matter {idx}: {matter.get('name', '') or matter.get('title', 'Untitled')}", level=3)
                
                p = doc.add_paragraph()
                p.add_run('Client: ').bold = True
                p.add_run(f"{matter.get('client', 'N/A')}\n")
                
                p.add_run('Status: ').bold = True
                p.add_run(f"{matter.get('publish_status', 'non_publishable')}\n")
                
                doc.add_paragraph().add_run('Matter Summary (Non-Publishable):').bold = True
                doc.add_paragraph(matter.get('optimizedText') or matter.get('optimized_text') or matter.get('description') or matter.get('rawNotes') or '')
                doc.add_paragraph("_" * 50)
        
    # Save the document
    file_path = f"{output_filename}.docx"
    doc.save(file_path)
    return file_path

