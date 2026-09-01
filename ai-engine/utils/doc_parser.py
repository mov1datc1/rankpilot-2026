import fitz  # PyMuPDF
from docx import Document
import os
import re
import hashlib
import tempfile
import urllib.request
from urllib.parse import urlparse
from collections import Counter

class DocumentParser:
    """
    Utility class to handle multi-format document ingestion (PDF & DOCX).
    Ensures text is extracted cleanly for LLM processing.
    v14.0: Added Trust Layer — programmatic matter counting and document stats.
    """
    
    @staticmethod
    def _collapse_exact_repetition(value: str) -> str:
        """Collapse SDT text duplicated two to four times by Word XML traversal."""

        candidate = (value or "").strip()
        for repetitions in (4, 3, 2):
            if len(candidate) % repetitions == 0:
                unit = candidate[: len(candidate) // repetitions]
                if unit and unit * repetitions == candidate:
                    return unit.strip()
        return candidate

    @staticmethod
    def parse(file_path: str) -> str:
        is_url = file_path.startswith('http://') or file_path.startswith('https://')
        
        # Parse extension correctly even with URL parameters
        parsed_path = urlparse(file_path).path if is_url else file_path
        extension = os.path.splitext(parsed_path)[1].lower()
        
        if extension not in ['.docx', '.doc', '.pdf']:
            raise ValueError(f"Unsupported file format: {extension}")

        local_path = file_path
        temp_file = None

        if is_url:
            # Download to a temporary file
            fd, local_path = tempfile.mkstemp(suffix=extension)
            os.close(fd)
            req = urllib.request.Request(file_path, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(local_path, 'wb') as out_file:
                out_file.write(response.read())
            temp_file = local_path

        try:
            if extension == '.docx':
                return DocumentParser._parse_docx(local_path)
            elif extension == '.doc':
                return DocumentParser._parse_doc(local_path)
            elif extension == '.pdf':
                return DocumentParser._parse_pdf(local_path)
        finally:
            # Cleanup temp file if it was created
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)

    @staticmethod
    def _parse_doc(file_path: str) -> str:
        """v24.3: Native Word 97-2003 (.doc) binary OLE extractor with LibreOffice conversion fallback."""
        # Method 1: Try libreoffice / soffice conversion if installed on system
        try:
            output_dir = tempfile.mkdtemp()
            docx_path = os.path.join(output_dir, os.path.splitext(os.path.basename(file_path))[0] + '.docx')
            ret = os.system(f'soffice --headless --convert-to docx "{file_path}" --outdir "{output_dir}" >/dev/null 2>&1')
            if ret == 0 and os.path.exists(docx_path):
                parsed = DocumentParser._parse_docx(docx_path)
                try:
                    os.remove(docx_path)
                    os.rmdir(output_dir)
                except Exception:
                    pass
                if len(parsed.strip()) > 100:
                    return parsed
        except Exception:
            pass

        # Method 2: Pure Python OLE Stream Text Extractor (No external dependencies)
        try:
            with open(file_path, 'rb') as f:
                content = f.read()

            lines = []
            # Extract UTF-16LE text strings
            try:
                text_utf16 = content.decode('utf-16le', errors='ignore')
                clean_utf16 = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text_utf16)
                blocks = re.findall(r'[\x20-\x7E\u00A0-\u024F\u1E00-\u1EFF]{4,}', clean_utf16)
                for b in blocks:
                    st = b.strip()
                    if len(st) > 3 and not any(st.startswith(x) for x in ['Root', 'WordDocument', 'þÿ', 'bjbj', 'Table', 'CompObj']):
                        if not any(c in st for c in ['Ą', 'ȫ']):
                            lines.append(st)
            except Exception:
                pass

            # Extract Latin1 text strings
            text_latin1 = content.decode('latin1', errors='ignore')
            blocks_latin = re.findall(r'[\x20-\x7E\xA0-\xFF]{4,}', text_latin1)
            for b in blocks_latin:
                st = b.strip()
                if len(st) > 4 and not any(st.startswith(x) for x in ['Root', 'WordDocument', 'þÿ', 'bjbj', 'Table', 'CompObj']):
                    if st not in lines:
                        lines.append(st)

            extracted_doc_text = '\n'.join(lines)
            if len(extracted_doc_text.strip()) > 50:
                return extracted_doc_text
        except Exception as doc_err:
            print(f"[DOC PARSER ERROR] OLE binary .doc extraction failed: {doc_err}")

        return ""

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Extracts text from Word documents in exact document order, handling SDT content controls and tables."""
        doc = Document(file_path)
        body = doc._body._element
        text_lines = []
        word_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        def xml_text(elem) -> str:
            # ``itertext()`` repeats content from nested Word SDT/run wrappers.
            # Reading only w:t leaves returns the text exactly once.
            return ''.join(node.text or '' for node in elem.findall(f'.//{word_ns}t'))

        def clean_text(raw: str) -> str:
            if not raw:
                return ''
            parts = [p.strip() for p in raw.split('\n') if p.strip()]
            cleaned_parts = []
            for p in parts:
                p = DocumentParser._collapse_exact_repetition(p)
                if not cleaned_parts or p != cleaned_parts[-1]:
                    cleaned_parts.append(p)
            return ' '.join(cleaned_parts)

        def process_node(elem):
            for child in elem:
                tag = child.tag.split('}')[-1]
                if tag == 'p':
                    p_txt = clean_text(xml_text(child).strip())
                    if p_txt:
                        text_lines.append(p_txt)
                elif tag == 'tbl':
                    for row in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                        row_cells = []
                        for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                            c_txt = clean_text(xml_text(cell).strip())
                            if c_txt and (not row_cells or c_txt != row_cells[-1]):
                                row_cells.append(c_txt)
                        if row_cells:
                            text_lines.append(' | '.join(row_cells))
                elif tag == 'sdt':
                    sdt_content = child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent')
                    if sdt_content is not None:
                        process_node(sdt_content)
                    else:
                        process_node(child)

        try:
            process_node(body)
        except Exception as err:
            print(f"[DOC PARSER WARNING] XML document-order traversal failed ({err}) — falling back to standard extraction")

        # Fallback if XML traversal yielded no text lines
        if not text_lines:
            for para in doc.paragraphs:
                if para.text.strip():
                    text_lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_lines.append(" | ".join(row_text))

        return "\n".join(text_lines)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Extracts text from PDF using PyMuPDF."""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text("text") + "\n"
        return text

    # =====================================================
    # v14.0 TRUST LAYER — Programmatic Source Verification
    # Rule 70: Extraction Validator
    # Rule 71: Pipeline Manifest
    # =====================================================

    MATTER_HEADER_PATTERN = re.compile(
        r'^\s*(Publishable|Confidential|Non[- ]publishable)\s+Matter\s+(\d+)\s*$',
        re.IGNORECASE,
    )

    @staticmethod
    def _canonical_matter_label(kind: str, number: int) -> str:
        normalized = (kind or "").casefold().replace(" ", "-")
        if normalized == "publishable":
            prefix = "Publishable"
        elif normalized == "confidential":
            prefix = "Confidential"
        else:
            prefix = "Non-publishable"
        return f"{prefix} Matter {int(number)}"

    @staticmethod
    def validate_matter_labels(labels: list) -> dict:
        """Validate uniqueness, numbering and section order for source headings."""

        normalized = []
        errors = []
        for raw in labels or []:
            match = DocumentParser.MATTER_HEADER_PATTERN.fullmatch(str(raw).strip())
            if not match:
                continue
            normalized.append(
                DocumentParser._canonical_matter_label(match.group(1), int(match.group(2)))
            )

        duplicate_labels = sorted(
            label for label, count in Counter(normalized).items() if count > 1
        )
        if duplicate_labels:
            errors.append(
                "Duplicate standalone matter headings: " + ", ".join(duplicate_labels)
            )

        section_numbers = {"Publishable": [], "Confidential": []}
        section_order = []
        for label in normalized:
            match = DocumentParser.MATTER_HEADER_PATTERN.fullmatch(label)
            if not match:
                continue
            kind = "Publishable" if match.group(1).casefold() == "publishable" else "Confidential"
            section_numbers[kind].append(int(match.group(2)))
            section_order.append(kind)

        for kind, numbers in section_numbers.items():
            if not numbers:
                continue
            unique_numbers = list(dict.fromkeys(numbers))
            expected = list(range(1, len(unique_numbers) + 1))
            if sorted(unique_numbers) != expected:
                errors.append(
                    f"{kind} matter headings must be contiguous from 1: "
                    f"found {sorted(unique_numbers)}, expected {expected}"
                )

        if "Confidential" in section_order:
            first_confidential = section_order.index("Confidential")
            if "Publishable" in section_order[first_confidential + 1:]:
                errors.append("Publishable matter headings appear after the confidential section")

        return {
            "passed": not errors,
            "errors": errors,
            "duplicate_labels": duplicate_labels,
            "normalized_labels": normalized,
        }

    @staticmethod
    def detect_rankpilot_generated_source(
        text: str, file_path: str = "", local_file_path: str = ""
    ) -> dict:
        """Reject RankPilot outputs as source evidence for another pipeline run."""

        reasons = []
        source = text or ""
        source_lower = source.casefold()
        explicit_markers = (
            "rankpilot — strategic audit letter",
            "pipeline manifest — trust layer",
            "canonical evidence reconciliation:",
            "rankpilot-generated-output",
        )
        for marker in explicit_markers:
            if marker in source_lower:
                reasons.append(f"RankPilot output marker found: {marker}")

        source_name = os.path.basename(urlparse(str(file_path)).path).casefold()
        if source_name.startswith(("rankpilot_submission_form_", "rankpilot_strategic_audit_")):
            reasons.append("RankPilot-generated output filename detected")

        # Legacy generated submissions predate the embedded provenance marker.
        # The combination below is emitted by RankPilot's generated lawyer table
        # and does not occur in the firm's original Chambers form.
        if (
            "current ranking:" in source_lower
            and "suggested ranking: suggested for ranking" in source_lower
        ):
            reasons.append("Legacy RankPilot-generated lawyer table detected")

        provenance_path = local_file_path or file_path
        if (
            provenance_path
            and str(provenance_path).lower().endswith(".docx")
            and os.path.exists(provenance_path)
        ):
            try:
                doc = Document(provenance_path)
                keywords = (doc.core_properties.keywords or "").casefold()
                if "rankpilot-generated-output" in keywords:
                    reasons.append("DOCX provenance identifies a RankPilot-generated output")
            except Exception:
                pass

        return {"passed": not reasons, "errors": reasons, "is_generated_output": bool(reasons)}

    @staticmethod
    def extract_matter_fields(section_text: str) -> dict:
        """Recover source D/E fields from one exact numbered matter section."""

        normalized = re.sub(r"\s+\|\s+", "\n", section_text or "")
        field_pattern = re.compile(
            r"(?ims)^\s*[DE]([1-9])\b[^\n]*\n(.*?)(?=^\s*[DE][1-9]\b|\Z)"
        )
        fields = {}
        for match in field_pattern.finditer(normalized):
            value = match.group(2).strip()
            value = re.sub(r"(?im)^\s*IMPORTANT:.*$", "", value).strip()
            instruction_patterns = (
                r"(?i)^this will be publishable\b.*$",
                r"(?i)^if you cannot reveal the client name\b.*$",
                r"(?i)^please say why this matter was important\b.*$",
                r"(?i)^also, tell us exactly what role\b.*$",
                r"(?i)^include currency and amount in figures\b.*$",
                r"(?i)^e\.g\.\s*link to press coverage\b.*$",
            )
            clean_lines = [
                line for line in value.splitlines()
                if line.strip()
                and not any(re.match(pattern, line.strip()) for pattern in instruction_patterns)
            ]
            value = "\n".join(clean_lines).strip()
            fields[int(match.group(1))] = value
        return {
            "client": fields.get(1, ""),
            "summary": fields.get(2, ""),
            "matter_value": fields.get(3, ""),
            "cross_border_jurisdictions": fields.get(4, ""),
            "lead_partner": fields.get(5, ""),
            "team_members": fields.get(6, ""),
            "other_firms": fields.get(7, ""),
            "completion_date": fields.get(8, ""),
        }

    @staticmethod
    def _count_matter_labels_in_text(text: str) -> dict:
        """Count numbered Chambers matter labels in normalized document text.

        This is the deterministic fallback for legacy ``.doc`` files after the
        parser has normalized their text.  Labels are de-duplicated by section
        and number, so repeated headers cannot inflate the manifest.
        """
        labels = []
        for line in (text or "").splitlines():
            match = DocumentParser.MATTER_HEADER_PATTERN.fullmatch(line.strip())
            if match:
                labels.append(
                    DocumentParser._canonical_matter_label(match.group(1), int(match.group(2)))
                )

        publishable = sum(label.startswith("Publishable") for label in labels)
        confidential = len(labels) - publishable
        result = {
            "total": len(labels),
            "publishable": publishable,
            "confidential": confidential,
            "matter_labels": labels,
        }
        result["label_validation"] = DocumentParser.validate_matter_labels(labels)
        return result

    @staticmethod
    def extract_numbered_matter_sections(text: str) -> dict:
        """Return verbatim matter sections keyed by normalized source label.

        The boundary is the next numbered matter header.  This gives the
        evidence ledger source text that predates all LLM transformation.
        DOCX table flattening can place a heading before a pipe or concatenate
        it directly with the first D/E field (``Matter 7D1``); both are still
        physical headings because they occur at the start of a parsed line.
        """
        pattern = re.compile(
            r'(?im)^\s*(Publishable|Confidential|Non[- ]publishable)\s+Matter\s+(\d+)'
            r'(?=\s*$|\s*\||[DE][1-9]\b)'
        )
        matches = list(pattern.finditer(text or ""))
        sections = {}
        for index, match in enumerate(matches):
            kind_raw = match.group(1).lower().replace(" ", "-")
            if kind_raw == "publishable":
                kind = "Publishable"
            elif kind_raw == "confidential":
                kind = "Confidential"
            else:
                kind = "Non-publishable"
            label = f"{kind} Matter {int(match.group(2))}"
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            excerpt = text[match.end():end].strip()
            sections[label.lower()] = {
                "label": label,
                "text": excerpt,
            }
        return sections

    @staticmethod
    def extract_c2_source(text: str) -> str:
        """Extract the submitted C2 answer without creating content for a blank field."""

        header = re.search(r"(?im)^\s*C2\b(?P<line>[^\n]*)", text or "")
        if not header:
            return ""
        inline = header.group("line").strip()
        if "|" in inline:
            inline = inline.rsplit("|", 1)[-1].strip()
        else:
            # A normal paragraph header contains only the question. Answers in
            # tabular extraction are separated with ``|`` and handled above.
            inline = ""
        remainder = (text or "")[header.end():]
        boundary = re.search(
            r"(?im)^\s*(?:C3\b|D\d*\b|D\.\s|WORK\s+HIGHLIGHTS\s+AND\s+CLIENTS|Publishable\s+Matter\s+\d+|Confidential\s+Matter\s+\d+)",
            remainder,
        )
        following = remainder[:boundary.start()] if boundary else remainder[:4000]
        content = "\n".join(part for part in (inline, following.strip()) if part).strip()
        for repetitions in (4, 3, 2):
            if len(content) % repetitions == 0:
                unit_length = len(content) // repetitions
                unit = content[:unit_length]
                if unit and unit * repetitions == content:
                    content = unit.strip()
                    break
        # Template instructions are not evidence. A section containing only an
        # instruction is treated as blank and therefore triggers a question.
        if re.fullmatch(
            r"(?is)(?:please\s+)?(?:provide|include|address|give).{0,300}", content
        ):
            return ""
        return content

    @staticmethod
    def extract_lawyer_roster(text: str) -> list:
        """Recover the complete B9 lawyer roster independently of the LLM.

        Legacy ``.doc`` extraction loses table columns but preserves the ordered
        names and profile URLs. A Chambers profile URL is treated as source
        evidence that the submitted lawyer is ranked; an internal firm profile
        alone is not converted into a ranking.
        """

        source = text or ""
        start_match = re.search(
            r"(?i)information regarding ranked and unranked lawyers", source
        )
        end_match = re.search(r"(?im)^\s*B10\b", source)
        if not start_match or not end_match or end_match.start() <= start_match.end():
            return []
        section = source[start_match.end():end_match.start()]
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        name_pattern = re.compile(
            r"^[A-ZÁÉÍÓÚÜÑ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ'’.-]*(?:\s+(?:[A-ZÁÉÍÓÚÜÑ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ'’.-]*|[A-Z])){1,5}$"
        )
        excluded = {
            "comments or web link", "partner ranked", "web link partner",
        }
        candidates = []
        for index, line in enumerate(lines):
            normalized = line.casefold()
            if (
                len(line) <= 80
                and name_pattern.fullmatch(line)
                and normalized not in excluded
                and not normalized.startswith(("current or", "please do", "comments or"))
            ):
                candidates.append((index, line))

        roster = []
        claimed_chambers_slugs = set()
        for position, (line_index, name) in enumerate(candidates):
            next_index = candidates[position + 1][0] if position + 1 < len(candidates) else len(lines)
            evidence_block = "\n".join(lines[line_index + 1:next_index]).casefold()
            chamber_slugs = re.findall(
                r"chambers\.com/lawyer/([a-z0-9-]+?)-latin-america",
                evidence_block,
                re.I,
            )
            if chamber_slugs:
                claimed_chambers_slugs.add(chamber_slugs[0].casefold())
            roster.append({
                "name": name,
                "is_partner": None,
                "is_ranked": "chambers.com/lawyer/" in evidence_block,
                "current_ranking": "Ranked" if "chambers.com/lawyer/" in evidence_block else None,
                "source_excerpt": "\n".join(lines[line_index:next_index]),
            })

        # Some binary DOC files lose a displayed name but retain the Chambers
        # URL. Recover that row from its slug only when no submitted name maps to it.
        seen_url_slugs = set()
        for slug in re.findall(r"chambers\.com/lawyer/([a-z0-9-]+?)-latin-america", section, re.I):
            normalized_slug = slug.casefold()
            if normalized_slug in seen_url_slugs or normalized_slug in claimed_chambers_slugs:
                continue
            seen_url_slugs.add(normalized_slug)
            display = " ".join(part.capitalize() for part in slug.split("-") if part)
            if display:
                roster.append({
                    "name": display,
                    "is_partner": None,
                    "is_ranked": True,
                    "current_ranking": "Ranked",
                    "source_excerpt": slug,
                })
        return roster

    @staticmethod
    def count_source_matters(file_path: str) -> dict:
        """
        Rule 70: Programmatically count matter sections. DOCX files are scanned
        through their XML tables; legacy DOC files use normalized text labels.
        This count is INDEPENDENT of LLM extraction and serves as ground truth.
        
        Returns:
            {
                "total": int,
                "publishable": int,
                "confidential": int,
                "matter_labels": ["Publishable Matter 1", "Confidential Matter 2", ...]
            }
        """
        is_url = file_path.startswith('http://') or file_path.startswith('https://')
        parsed_path = urlparse(file_path).path if is_url else file_path
        extension = os.path.splitext(parsed_path)[1].lower()
        
        if extension not in {'.docx', '.doc'}:
            return {"total": 0, "publishable": 0, "confidential": 0, "matter_labels": [], "note": "Programmatic count is available for DOC and DOCX only"}
        
        local_path = file_path
        temp_file = None
        
        if is_url:
            fd, local_path = tempfile.mkstemp(suffix=extension)
            os.close(fd)
            req = urllib.request.Request(file_path, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(local_path, 'wb') as out_file:
                out_file.write(response.read())
            temp_file = local_path
        
        try:
            if extension == '.doc':
                normalized_text = DocumentParser._parse_doc(local_path)
                result = DocumentParser._count_matter_labels_in_text(normalized_text)
                result["count_method"] = "normalized_doc_text"
                return result

            doc = Document(local_path)
            matter_labels = []
            seen_table_indices = set()
            
            # Only a standalone first paragraph/cell can open a matter table. A
            # label mentioned in a client, narrative or audit note is not a new
            # matter and must never inflate the source manifest.
            all_tbl_elems = doc._body._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')
            for ti, tbl_elem in enumerate(all_tbl_elems):
                if ti in seen_table_indices:
                    continue
                paragraphs = tbl_elem.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
                )
                heading = ""
                for paragraph in paragraphs:
                    paragraph_text = DocumentParser._collapse_exact_repetition(
                        ''.join(
                            node.text or ''
                            for node in paragraph.findall(
                                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
                            )
                        ).strip()
                    )
                    match = DocumentParser.MATTER_HEADER_PATTERN.fullmatch(paragraph_text)
                    if match:
                        heading = paragraph_text
                        break
                match = DocumentParser.MATTER_HEADER_PATTERN.fullmatch(heading)
                if match:
                    matter_labels.append(
                        DocumentParser._canonical_matter_label(
                            match.group(1), int(match.group(2))
                        )
                    )
                    seen_table_indices.add(ti)
            
            # Also scan paragraphs for matter headers (some templates use headings)
            for para in doc.paragraphs:
                txt = DocumentParser._collapse_exact_repetition(para.text.strip())
                match = DocumentParser.MATTER_HEADER_PATTERN.fullmatch(txt)
                if match:
                    label = DocumentParser._canonical_matter_label(
                        match.group(1), int(match.group(2))
                    )
                    matter_labels.append(label)
            
            # Classify
            publishable = sum(1 for l in matter_labels 
                            if 'publishable' in l.lower() 
                            and 'non' not in l.lower() 
                            and 'confidential' not in l.lower())
            confidential = sum(1 for l in matter_labels 
                             if 'confidential' in l.lower() 
                             or 'non-publishable' in l.lower() 
                             or 'non publishable' in l.lower())
            
            result = {
                "total": len(matter_labels),
                "publishable": publishable,
                "confidential": confidential,
                "matter_labels": matter_labels,
                "count_method": "docx_xml",
            }
            result["label_validation"] = DocumentParser.validate_matter_labels(matter_labels)
            return result
        except Exception as e:
            print(f"[MATTER COUNTER] Error counting matters: {e}")
            return {"total": 0, "publishable": 0, "confidential": 0, "matter_labels": [], "error": str(e)}
        finally:
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)

    @staticmethod
    def get_document_stats(file_path: str, source_file_name: str = "") -> dict:
        """
        Rule 71: Generate document-level statistics for the Pipeline Manifest.
        Provides ground truth about what the system actually read.
        
        Returns:
            {
                "file_name": str,
                "file_hash": str (SHA-256 of content),
                "word_count": int,
                "paragraph_count": int,
                "table_count": int,
                "source_matters": { ... from count_source_matters },
            }
        """
        is_url = file_path.startswith('http://') or file_path.startswith('https://')
        parsed_path = urlparse(file_path).path if is_url else file_path
        extension = os.path.splitext(parsed_path)[1].lower()
        file_name = source_file_name or os.path.basename(parsed_path)
        
        stats = {
            "file_name": file_name,
            "file_hash": "",
            "word_count": 0,
            "paragraph_count": 0,
            "table_count": 0,
            "source_matters": {},
        }
        
        local_path = file_path
        temp_file = None
        
        if is_url:
            fd, local_path = tempfile.mkstemp(suffix=extension)
            os.close(fd)
            req = urllib.request.Request(file_path, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(local_path, 'wb') as out_file:
                out_file.write(response.read())
            temp_file = local_path
        
        source_text = ""
        try:
            # File hash
            with open(local_path, 'rb') as f:
                stats["file_hash"] = hashlib.sha256(f.read()).hexdigest()
            
            if extension == '.docx':
                source_text = DocumentParser._parse_docx(local_path)
                stats["word_count"] = len(source_text.split())
                stats["paragraph_count"] = len([l for l in source_text.splitlines() if l.strip()])
                doc = Document(local_path)
                stats["table_count"] = len(doc._body._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'))
                stats["source_matters"] = DocumentParser.count_source_matters(local_path)
            elif extension == '.doc':
                source_text = DocumentParser._parse_doc(local_path)
                stats["word_count"] = len(source_text.split())
                stats["paragraph_count"] = len([line for line in source_text.splitlines() if line.strip()])
                stats["source_matters"] = DocumentParser.count_source_matters(local_path)
            elif extension == '.pdf':
                source_text = DocumentParser._parse_pdf(local_path)
                stats["word_count"] = len(source_text.split())
                stats["paragraph_count"] = source_text.count('\n\n') + 1

            label_validation = stats.get("source_matters", {}).get(
                "label_validation", {"passed": True, "errors": []}
            )
            provenance_validation = DocumentParser.detect_rankpilot_generated_source(
                source_text, source_file_name or file_path, local_path
            )
            errors = list(label_validation.get("errors", [])) + list(
                provenance_validation.get("errors", [])
            )
            stats["source_validation"] = {
                "passed": not errors,
                "errors": errors,
                "label_validation": label_validation,
                "provenance_validation": provenance_validation,
            }
            stats["source_format"] = extension.lstrip(".")
        except Exception as e:
            stats["error"] = str(e)
        finally:
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)
        
        return stats

    # =====================================================
    # v14.1 PRE-FLIGHT GATE — Rule 72: Ranking Evidence Detector
    # Scans DOCX for evidence that the firm is already ranked.
    # Used to flag contradictions when user declares "Unranked".
    # =====================================================
    @staticmethod
    def detect_ranking_evidence(file_path: str) -> dict:
        """
        Rule 72: Scan DOCX for evidence of existing ranking status.
        Looks for:
        - C2/Feedback sections mentioning "current rankings", "our coverage"
        - B9 sections listing "Ranked" lawyers
        - Any text mentioning "Band", "Tier", "currently ranked"
        
        Returns:
            {
                "has_ranking_evidence": bool,
                "evidence_type": str,     # "explicit" | "implicit" | "none"
                "evidence_text": str,     # The key phrase found
                "detected_band": str,     # If a specific band/tier is mentioned
                "ranked_lawyers": list,   # Names flagged as currently ranked
            }
        """
        is_url = file_path.startswith('http://') or file_path.startswith('https://')
        parsed_path = urlparse(file_path).path if is_url else file_path
        extension = os.path.splitext(parsed_path)[1].lower()
        
        result = {
            "has_ranking_evidence": False,
            "evidence_type": "none",
            "evidence_text": "",
            "detected_band": "",
            "ranked_lawyers": [],
        }
        
        if extension != '.docx':
            return result
        
        local_path = file_path
        temp_file = None
        
        if is_url:
            fd, local_path = tempfile.mkstemp(suffix=extension)
            os.close(fd)
            req = urllib.request.Request(file_path, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(local_path, 'wb') as out_file:
                out_file.write(response.read())
            temp_file = local_path
        
        try:
            doc = Document(local_path)
            all_text_blocks = []
            
            # Collect text from tables for scanning
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt and len(txt) > 20:
                            all_text_blocks.append((f"T{ti}R{ri}", txt))
            
            # Also scan paragraphs
            for pi, para in enumerate(doc.paragraphs):
                txt = para.text.strip()
                if txt and len(txt) > 20:
                    all_text_blocks.append((f"P{pi}", txt))
            
            # Pattern 1: EXPLICIT ranking evidence
            # "current rankings", "our coverage", "currently ranked", "Band N"
            explicit_patterns = [
                (r'current\s+rank(?:ing|s)', 'explicit'),
                (r'our\s+(?:current\s+)?coverage', 'explicit'),
                (r'currently\s+(?:ranked|listed|included)', 'explicit'),
                (r'(?:Band|Tier)\s+\d', 'explicit'),
                (r'maintain\s+(?:our|the)\s+(?:current\s+)?rank', 'explicit'),
                (r'we\s+(?:are|remain)\s+ranked', 'explicit'),
            ]
            
            for location, text in all_text_blocks:
                text_lower = text.lower()
                for pattern, ev_type in explicit_patterns:
                    match = re.search(pattern, text_lower)
                    if match:
                        result["has_ranking_evidence"] = True
                        result["evidence_type"] = ev_type
                        # Extract surrounding context (max 200 chars)
                        start = max(0, match.start() - 50)
                        end = min(len(text), match.end() + 100)
                        result["evidence_text"] = text[start:end].strip()
                        
                        # Try to extract specific band
                        band_match = re.search(r'(?:Band|Tier)\s+(\d)', text, re.IGNORECASE)
                        if band_match:
                            result["detected_band"] = f"Band {band_match.group(1)}"
                        break
                if result["has_ranking_evidence"]:
                    break
            
            # Pattern 2: IMPLICIT evidence — check B9/lawyer table for "Ranked" column
            for ti, table in enumerate(doc.tables):
                if not table.rows:
                    continue
                header_text = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
                if 'ranked' in header_text and ('lawyer' in header_text or 'unranked' in header_text):
                    # This is a B9-style table. Check for ranked lawyers
                    for ri in range(1, len(table.rows)):
                        row = table.rows[ri]
                        cells = [c.text.strip() for c in row.cells]
                        name = cells[0] if cells else ""
                        # Check if any cell says "Ranked", "Band", "Star" etc.
                        row_text = " ".join(cells).lower()
                        if any(k in row_text for k in ['ranked', 'band', 'star', 'tier']):
                            if name and name.lower() not in ['name', '']:
                                result["ranked_lawyers"].append(name)
                    
                    if result["ranked_lawyers"]:
                        result["has_ranking_evidence"] = True
                        if result["evidence_type"] == "none":
                            result["evidence_type"] = "implicit"
                            result["evidence_text"] = f"B9 table lists {len(result['ranked_lawyers'])} lawyers with ranking indicators"
            
        except Exception as e:
            result["error"] = str(e)
        finally:
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)
        
        return result

    # =====================================================
    # v14.1 PRE-FLIGHT GATE — Rule 73: Directory Template Detector
    # Auto-detects directory format, practice area, jurisdiction
    # from the DOCX structure itself — independent of user input.
    # =====================================================
    @staticmethod
    def detect_directory_template(file_path: str) -> dict:
        """
        Rule 73: Auto-detect directory, practice area, jurisdiction, and firm name
        from the DOCX template structure.
        
        Chambers: Has A1 (Firm name), A2 (Practice Area), A3 (Location) tables
        Legal 500: Different structure — firm name in first table, practice area 
                   in department area, matter tables use different patterns
        
        Returns:
            {
                "detected_directory": str,    # "Chambers" | "Legal 500" | "IFLR1000" | "Unknown"
                "detected_firm_name": str,
                "detected_practice_area": str,
                "detected_jurisdiction": str,
                "confidence": str,            # "high" | "medium" | "low"
                "detection_signals": list,    # What signals were used
            }
        """
        is_url = file_path.startswith('http://') or file_path.startswith('https://')
        parsed_path = urlparse(file_path).path if is_url else file_path
        extension = os.path.splitext(parsed_path)[1].lower()
        
        result = {
            "detected_directory": "Unknown",
            "detected_firm_name": "",
            "detected_practice_area": "",
            "detected_jurisdiction": "",
            "confidence": "low",
            "detection_signals": [],
        }
        
        if extension != '.docx':
            return result
        
        local_path = file_path
        temp_file = None
        
        if is_url:
            fd, local_path = tempfile.mkstemp(suffix=extension)
            os.close(fd)
            req = urllib.request.Request(file_path, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(local_path, 'wb') as out_file:
                out_file.write(response.read())
            temp_file = local_path
        
        try:
            doc = Document(local_path)
            
            if len(doc.tables) < 3:
                return result
            
            # === CHAMBERS DETECTION ===
            # Chambers templates have: T0 = "A1 Firm name", T1 = "A2 Practice Area", T2 = "A3 Location"
            t0_header = doc.tables[0].rows[0].cells[0].text.strip().lower() if doc.tables[0].rows else ""
            t1_header = doc.tables[1].rows[0].cells[0].text.strip().lower() if len(doc.tables) > 1 and doc.tables[1].rows else ""
            t2_header = doc.tables[2].rows[0].cells[0].text.strip().lower() if len(doc.tables) > 2 and doc.tables[2].rows else ""
            
            is_chambers = (
                ('a1' in t0_header and 'firm' in t0_header) or
                ('a2' in t1_header and 'practice' in t1_header) or
                ('a3' in t2_header and 'location' in t2_header)
            )
            
            if is_chambers:
                result["detected_directory"] = "Chambers"
                result["confidence"] = "high"
                result["detection_signals"].append("Chambers A1/A2/A3 table structure detected")
                
                # Extract firm name (T0 R1)
                if len(doc.tables[0].rows) > 1:
                    result["detected_firm_name"] = doc.tables[0].rows[1].cells[0].text.strip()
                
                # Extract practice area (T1 R1)
                if len(doc.tables) > 1 and len(doc.tables[1].rows) > 1:
                    result["detected_practice_area"] = doc.tables[1].rows[1].cells[0].text.strip()
                
                # Extract jurisdiction (T2 R1)
                if len(doc.tables) > 2 and len(doc.tables[2].rows) > 1:
                    result["detected_jurisdiction"] = doc.tables[2].rows[1].cells[0].text.strip()
                
                return result
            
            # === LEGAL 500 DETECTION ===
            # Legal 500 templates: T0 = firm name (1 row), matter tables have "Publishable matter" / "Non-publishable matter"
            has_legal500_matters = False
            for ti, table in enumerate(doc.tables):
                if table.rows:
                    first_cell = table.rows[0].cells[0].text.strip().lower()
                    if re.match(r'^(?:publishable|non-publishable)\s+matter', first_cell):
                        has_legal500_matters = True
                        result["detection_signals"].append(f"Legal 500 matter format at table {ti}")
                        break
            
            # Legal 500 also has "leading partner 1", "next generation partner 1" etc.
            has_legal500_lawyers = False
            for ti, table in enumerate(doc.tables):
                if table.rows:
                    first_cell = table.rows[0].cells[0].text.strip().lower()
                    if 'leading partner' in first_cell or 'next generation partner' in first_cell:
                        has_legal500_lawyers = True
                        result["detection_signals"].append(f"Legal 500 lawyer format at table {ti}")
                        break
            
            if has_legal500_matters or has_legal500_lawyers:
                result["detected_directory"] = "Legal 500"
                result["confidence"] = "high" if (has_legal500_matters and has_legal500_lawyers) else "medium"
                
                # Firm name — usually T0 R0 (single-row table)
                if len(doc.tables[0].rows) == 1:
                    result["detected_firm_name"] = doc.tables[0].rows[0].cells[0].text.strip()
                
                # Practice area — try to find in file name or in known tables
                # Legal 500 doesn't have a clean "A2 Practice Area" — scan for it
                file_name = os.path.basename(parsed_path).lower()
                for pa in ['labour', 'employment', 'banking', 'finance', 'corporate', 'disputes', 
                          'tax', 'real estate', 'intellectual property', 'energy', 'competition']:
                    if pa in file_name:
                        result["detected_practice_area"] = pa.title()
                        result["detection_signals"].append(f"Practice area '{pa}' detected from filename")
                        break
                
                # Jurisdiction from lawyer locations
                for ti, table in enumerate(doc.tables):
                    if table.rows and len(table.rows) > 1:
                        header = table.rows[0].cells[0].text.strip().lower()
                        if header == 'name' and len(table.rows[0].cells) > 1:
                            loc_header = table.rows[0].cells[-1].text.strip().lower()
                            if 'location' in loc_header:
                                for ri in range(1, len(table.rows)):
                                    loc = table.rows[ri].cells[-1].text.strip()
                                    if loc and loc.lower() not in ['location', '']:
                                        result["detected_jurisdiction"] = loc
                                        break
                                break
                
                return result
            
            # === FALLBACK: Try filename ===
            file_name_lower = os.path.basename(parsed_path).lower()
            if 'chambers' in file_name_lower:
                result["detected_directory"] = "Chambers"
                result["confidence"] = "medium"
                result["detection_signals"].append("Directory 'Chambers' detected from filename")
            elif 'legal 500' in file_name_lower or 'legal500' in file_name_lower:
                result["detected_directory"] = "Legal 500"
                result["confidence"] = "medium"
                result["detection_signals"].append("Directory 'Legal 500' detected from filename")
            elif 'iflr' in file_name_lower:
                result["detected_directory"] = "IFLR1000"
                result["confidence"] = "medium"
                result["detection_signals"].append("Directory 'IFLR1000' detected from filename")
            
        except Exception as e:
            result["error"] = str(e)
        finally:
            if temp_file and os.path.exists(temp_file):
                os.remove(temp_file)
        
        return result
