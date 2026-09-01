import sys
import unittest
import os
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile


AI_ENGINE = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(AI_ENGINE))

from utils.ooxml_validation import validate_docx_ooxml  # noqa: E402
from core.docx_cloner import _normalize_docx_table_widths, clone_and_replace  # noqa: E402
from utils.doc_parser import DocumentParser  # noqa: E402
from docx import Document  # noqa: E402


def package(document_xml: str) -> bytes:
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", document_xml)
    return stream.getvalue()


class OoxmlValidationTests(unittest.TestCase):
    def test_dxa_table_passes(self):
        xml = """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:tbl><w:tblPr><w:tblW w:w="9360" w:type="dxa"/></w:tblPr><w:tblGrid><w:gridCol w:w="4680"/><w:gridCol w:w="4680"/></w:tblGrid></w:tbl></w:body></w:document>"""
        self.assertEqual([], validate_docx_ooxml(package(xml)))

    def test_percentage_and_zero_grid_fail(self):
        xml = """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:tbl><w:tblPr><w:tblW w:w="100" w:type="pct"/></w:tblPr><w:tblGrid><w:gridCol w:w="0"/></w:tblGrid></w:tbl></w:body></w:document>"""
        errors = validate_docx_ooxml(package(xml))
        self.assertTrue(any("non-DXA" in error for error in errors))
        self.assertTrue(any("invalid width" in error for error in errors))

    def test_cloner_normalizes_all_tables_to_dxa(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Left"
        table.cell(0, 1).text = "Right"
        _normalize_docx_table_widths(document)
        stream = BytesIO()
        document.save(stream)
        self.assertEqual([], validate_docx_ooxml(stream.getvalue()))

    def test_docx_counter_ignores_inline_matter_reference(self):
        document = Document()
        table = document.add_table(rows=3, cols=1)
        table.cell(0, 0).text = "Publishable Matter 1"
        table.cell(1, 0).text = "D1 Name of client"
        table.cell(2, 0).text = "Not stated in source (Confidential Matter 2)"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp:
            path = temp.name
        try:
            document.save(path)
            result = DocumentParser.count_source_matters(path)
            self.assertEqual(["Publishable Matter 1"], result["matter_labels"])
        finally:
            os.remove(path)

    def test_cloner_changes_summary_only_and_adds_provenance(self):
        document = Document()
        table = document.add_table(rows=8, cols=1)
        values = [
            "D1 Name of client", "Client Alpha",
            "D2 Summary of matter and your department's involvement", "Original summary",
            "D3 Value of deal / matter", "US$10 million",
            "D5 Lead partner", "Lawyer One",
        ]
        for index, value in enumerate(values):
            table.cell(index, 0).text = value
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp:
            path = temp.name
        try:
            document.save(path)
            output = clone_and_replace(path, enhanced_matters=[{
                "client": "Client Alpha",
                "optimized_text": "Optimized source-bounded summary",
                "matter_value": "WRONG VALUE",
                "lead_partner": "WRONG LAWYER",
                "publish_status": "publishable",
            }])
            generated = Document(BytesIO(output))
            all_text = "\n".join(cell.text for row in generated.tables[0].rows for cell in row.cells)
            self.assertIn("Optimized source-bounded summary", all_text)
            self.assertIn("US$10 million", all_text)
            self.assertIn("Lawyer One", all_text)
            self.assertNotIn("WRONG VALUE", all_text)
            self.assertNotIn("WRONG LAWYER", all_text)
            self.assertIn("rankpilot-generated-output", generated.core_properties.keywords)
        finally:
            os.remove(path)

    def test_cloner_fails_when_matter_counts_do_not_match(self):
        document = Document()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp:
            path = temp.name
        try:
            document.save(path)
            with self.assertRaisesRegex(ValueError, "Matter table count mismatch"):
                clone_and_replace(path, enhanced_matters=[{
                    "client": "Unsupported Client",
                    "optimized_text": "Text",
                    "publish_status": "publishable",
                }])
        finally:
            os.remove(path)


if __name__ == "__main__":
    unittest.main()
